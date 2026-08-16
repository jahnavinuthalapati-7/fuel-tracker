import streamlit as st
from supabase import create_client, Client
import pandas as pd
from datetime import datetime
import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from docx import Document

# ============ SUPABASE CONFIG ============
SUPABASE_URL = "https://yccutkrmflxapwtjngep.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InljY3V0a3JtZmx4YXB3dGpuZ2VwIiwicm9sZSI6ImFub24iLCJpYXQiOjE3ODQ4NDM5MzEsImV4cCI6MjEwMDQxOTkzMX0.DsWCR0tYq895So5oJWD7Jwia_HRVxO09Y9rv2_Wns9w"
supabase_client: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# ============ PAGE CONFIG ============
st.set_page_config(
    page_title="Fuel Tracker",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={'Get Help': None, 'Report a bug': None, 'About': "Fuel Tracking System"}
)

# ============ MINIMAL CSS - ONLY FOR LAYOUT ============
st.markdown("""
<style>
    .main { padding: 20px; }
    .metric-box {
        background-color: #f5f5f5;
        padding: 20px;
        border-radius: 10px;
        border-left: 4px solid #667eea;
    }
</style>
""", unsafe_allow_html=True)

# ============ SESSION STATE ============
if 'user' not in st.session_state:
    st.session_state.user = None

# ============ EXPORT FUNCTIONS ============
def export_to_excel(data):
    wb = Workbook()
    ws = wb.active
    ws.title = "Fuel Entries"
    
    headers = ["Date", "Slip No", "Type", "Fuel Type", "Vehicle No", "VIN", "Reg No", "Model",
               "Allocated", "Opening Odo", "Closing Odo", "KM", "Qty", "Rate", "Amount",
               "Mileage", "Opening Stock", "Closing Stock", "Approved By", "Remarks"]
    ws.append(headers)
    
    header_fill = PatternFill(start_color="667eea", end_color="667eea", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
    
    for row in data:
        ws.append([
            row.get('date', ''),
            row.get('slip_no', ''),
            row.get('transaction_type', ''),
            row.get('fuel_type', ''),
            row.get('vehicle_no', ''),
            row.get('vin_no', ''),
            row.get('registration_no', ''),
            row.get('model_no', ''),
            row.get('allocated_to', ''),
            row.get('opening_odometer', ''),
            row.get('closing_odometer', ''),
            row.get('km_run', ''),
            row.get('quantity', '') or row.get('issue_quantity', ''),
            row.get('rate', ''),
            row.get('amount', ''),
            row.get('mileage', ''),
            row.get('opening_stock', ''),
            row.get('closing_stock', ''),
            row.get('approved_by', ''),
            row.get('remarks', '')
        ])
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()

def export_to_pdf(data):
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter)
    elements = []
    
    styles = getSampleStyleSheet()
    title = Paragraph("Fuel Tracker Report", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 12))
    
    table_data = [["Date", "Slip", "Type", "Fuel", "Vehicle", "Qty", "Amount", "Approved By"]]
    
    for row in data:
        table_data.append([
            str(row.get('date', '')),
            str(row.get('slip_no', '')),
            str(row.get('transaction_type', '')),
            str(row.get('fuel_type', '')),
            str(row.get('vehicle_no', '')),
            str(row.get('quantity', '') or row.get('issue_quantity', '')),
            str(row.get('amount', '')),
            str(row.get('approved_by', ''))
        ])
    
    table = Table(table_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), '#667eea'),
        ('TEXTCOLOR', (0, 0), (-1, 0), 'white'),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, 'black')
    ]))
    
    elements.append(table)
    doc.build(elements)
    output.seek(0)
    return output.getvalue()

def export_to_word(data):
    doc = Document()
    doc.add_heading('Fuel Tracker Report', 0)
    
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    headers = ["Date", "Slip No", "Type", "Fuel", "Qty", "Amount", "Vehicle", "Approved By"]
    
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    for row in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row.get('date', ''))
        row_cells[1].text = str(row.get('slip_no', ''))
        row_cells[2].text = str(row.get('transaction_type', ''))
        row_cells[3].text = str(row.get('fuel_type', ''))
        row_cells[4].text = str(row.get('quantity', '') or row.get('issue_quantity', ''))
        row_cells[5].text = str(row.get('amount', ''))
        row_cells[6].text = str(row.get('vehicle_no', ''))
        row_cells[7].text = str(row.get('approved_by', ''))
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

# ============ LOGIN PAGE ============
if st.session_state.user is None:
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.title("⛽ Fuel Tracker")
        st.markdown("### Fuel Entry & Management System")
        st.divider()
        
        username = st.text_input("Username", placeholder="Enter username")
        password = st.text_input("Password", type="password", placeholder="Enter password")
        
        if st.button("🔐 Login", use_container_width=True):
            try:
                response = supabase_client.table('users').select('*').eq('username', username).execute()
                
                if response.data and len(response.data) > 0:
                    user = response.data[0]
                    if user['password'] == password:
                        st.session_state.user = user
                        st.success("✅ Login successful!")
                        st.rerun()
                    else:
                        st.error("❌ Invalid credentials")
                else:
                    st.error("❌ Invalid credentials")
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")

# ============ MAIN APP ============
else:
    # SIDEBAR
    with st.sidebar:
        st.title("⛽ Fuel Tracker")
        st.markdown(f"👤 **{st.session_state.user['username']}**")
        st.markdown(f"🎯 {st.session_state.user['role']}")
        st.divider()
        
        if st.button("🔴 Logout", use_container_width=True):
            st.session_state.user = None
            st.rerun()
        
        st.divider()
        
        page = st.radio(
            "📋 Menu",
            ["📊 Dashboard", "✏️ Fuel Entry", "📈 Reports", "👥 Users", "🔑 Password"]
        )
    
    # ============ DASHBOARD PAGE ============
    if page == "📊 Dashboard":
        st.title("📊 Dashboard")
        st.write("Real-time stock tracking by fuel type")
        st.divider()
        
        try:
            entries = supabase_client.table('fuel_entries').select('*').execute()
            entries_data = entries.data if entries.data else []
            
            # Show metrics for each fuel type
            for fuel_type in ['Diesel', 'Petrol']:
                fuel_entries = [e for e in entries_data if e['fuel_type'] == fuel_type]
                old_stock = fuel_entries[0]['opening_stock'] if fuel_entries and fuel_entries[0].get('opening_stock') else 0
                purchased = sum([e.get('quantity', 0) or 0 for e in fuel_entries if e['transaction_type'] == 'Purchase'])
                allocated = sum([e.get('issue_quantity', 0) or 0 for e in fuel_entries if e['transaction_type'] == 'Issue'])
                total = old_stock + purchased - allocated
                
                st.subheader(f"💧 {fuel_type}")
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("OLD STOCK", f"{old_stock:.2f}L")
                with col2:
                    st.metric("PURCHASED", f"+{purchased:.2f}L")
                with col3:
                    st.metric("ALLOCATED", f"-{allocated:.2f}L")
                with col4:
                    st.metric("TOTAL", f"{total:.2f}L")
                
                st.divider()
            
            # Recent entries table
            st.subheader("📋 Recent Entries")
            search = st.text_input("🔍 Search Slip No, Vehicle, Fuel Type, Allocated To, Approved By...")
            
            if search:
                filtered = [e for e in entries_data if
                           search.lower() in str(e.get('slip_no', '')).lower() or
                           search.lower() in str(e.get('vehicle_no', '')).lower() or
                           search.lower() in str(e.get('fuel_type', '')).lower() or
                           search.lower() in str(e.get('allocated_to', '')).lower() or
                           search.lower() in str(e.get('approved_by', '')).lower()]
            else:
                filtered = entries_data[:10]
            
            if filtered:
                df = pd.DataFrame(filtered)
                display_cols = ['date', 'transaction_type', 'slip_no', 'vehicle_no', 'fuel_type', 'quantity', 'issue_quantity', 'approved_by']
                display_cols = [col for col in display_cols if col in df.columns]
                st.dataframe(df[display_cols], use_container_width=True, hide_index=True)
                
                # DELETE SECTION
                st.divider()
                st.subheader("🗑️ Delete Entries")
                delete_col1, delete_col2, delete_col3 = st.columns([2, 2, 1])
                with delete_col1:
                    slip_to_delete = st.selectbox("Select Slip No to delete:", df['slip_no'].unique(), key="delete_slip")
                with delete_col2:
                    entry_to_delete = st.selectbox("Transaction Type:", df['transaction_type'].unique(), key="delete_type")
                with delete_col3:
                    if st.button("🗑️ Delete", use_container_width=True, key="delete_btn"):
                        try:
                            entry_id = df[(df['slip_no'] == slip_to_delete) & (df['transaction_type'] == entry_to_delete)]['id'].values
                            if len(entry_id) > 0:
                                supabase_client.table('fuel_entries').delete().eq('id', int(entry_id[0])).execute()
                                st.success("✅ Entry deleted successfully!")
                                st.rerun()
                            else:
                                st.error("❌ Entry not found")
                        except Exception as e:
                            st.error(f"❌ Error deleting: {e}")
            else:
                st.info("📭 No entries found")
                
        except Exception as e:
            st.error(f"❌ Error: {e}")
    
    # ============ FUEL ENTRY PAGE ============
    elif page == "✏️ Fuel Entry":
        st.title("✏️ Add Fuel Entry")
        st.write("Record PURCHASE or ISSUE transactions")
        st.divider()
        
        tab1, tab2 = st.tabs(["🛢️ Purchase Fuel", "🚗 Issue to Vehicle"])
        
        with tab1:
            st.subheader("🛢️ Purchase Fuel")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                slip_no = st.text_input("Slip No *", key="slip_p")
            with col2:
                fuel_type = st.selectbox("Fuel Type *", ["Diesel", "Petrol"], key="fuel_p")
            with col3:
                stock_type = st.selectbox("Stock Type *", ["New Stock", "Old Stock"], key="st_p")
            
            opening_stock = None
            if stock_type == "Old Stock":
                opening_stock = st.number_input("Stock Amount (L)", min_value=0.0, step=0.01, key="os_p")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                quantity = st.number_input("Quantity (L) *", min_value=0.0, step=0.01, key="qty_p")
            with col2:
                rate = st.number_input("Rate (per L)", min_value=0.0, step=0.01, key="rate_p")
            with col3:
                amount = quantity * rate
                st.metric("Amount (₹)", f"{amount:.2f}")
            
            col1, col2 = st.columns(2)
            with col1:
                vendor_name = st.text_input("Vendor Name", key="vn_p")
            with col2:
                vendor_location = st.text_input("Vendor Location", key="vl_p")
            
            col1, col2 = st.columns(2)
            with col1:
                approved_by = st.text_input("Approved By", key="ab_p")
            with col2:
                remarks = st.text_input("Remarks", key="rm_p")
            
            if st.button("✅ Save Purchase", use_container_width=True, key="save_p"):
                if not slip_no or not fuel_type:
                    st.error("❌ Please fill required fields")
                else:
                    try:
                        supabase_client.table('fuel_entries').insert({
                            'slip_no': slip_no,
                            'transaction_type': 'Purchase',
                            'fuel_type': fuel_type,
                            'date': datetime.now().strftime("%Y-%m-%d"),
                            'opening_stock': opening_stock,
                            'quantity': quantity,
                            'rate': rate,
                            'amount': amount,
                            'vendor_name': vendor_name,
                            'vendor_location': vendor_location,
                            'approved_by': approved_by,
                            'remarks': remarks
                        }).execute()
                        st.success("✅ Purchase saved successfully!")
                        st.balloons()
                        st.rerun()
                    except Exception as e:
                        if "duplicate key" in str(e).lower():
                            st.warning("⚠️ This Slip No already exists as a Purchase. Try a different number.")
                        else:
                            st.error(f"Error: {str(e)}")
        
        with tab2:
            st.subheader("🚗 Issue to Vehicle")
            
            col1, col2 = st.columns(2)
            with col1:
                slip_no = st.text_input("Slip No *", key="slip_i")
            with col2:
                fuel_type = st.selectbox("Fuel Type *", ["Diesel", "Petrol"], key="fuel_i")
            
            vehicle_no = st.text_input("Vehicle No *", key="vehicle_i")
            vin_no = st.text_input("VIN No", key="vin_i")
            registration_no = st.text_input("Registration No", key="reg_i")
            model_no = st.text_input("Model No", key="model_i")
            allocated_to = st.selectbox("Allocated To", ["Demo Car", "Assigned Car", "Customer Car", "Loaner Car", "Other"], key="alloc_i")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                opening_odo = st.number_input("Opening Odometer", min_value=0.0, step=0.01, key="odo_o_i")
            with col2:
                closing_odo = st.number_input("Closing Odometer", min_value=0.0, step=0.01, key="odo_c_i")
            with col3:
                km_run = closing_odo - opening_odo
                st.metric("KM Run", f"{km_run:.2f}")
            
            col1, col2 = st.columns(2)
            with col1:
                issue_quantity = st.number_input("Issue Quantity (L) *", min_value=0.0, step=0.01, key="iq_i")
            with col2:
                mileage = km_run / issue_quantity if issue_quantity > 0 else 0
                st.metric("Mileage (KM/L)", f"{mileage:.2f}")
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                recipient_name = st.text_input("Recipient Name", key="rec_i")
            with col2:
                receiver_name = st.text_input("Receiver Name", key="recv_i")
            with col3:
                approved_by = st.text_input("Approved By", key="ab_i")
            with col4:
                remarks = st.text_input("Remarks", key="rm_i")
            
            if st.button("✅ Save Issue", use_container_width=True, key="save_i"):
                if not slip_no or not fuel_type or not vehicle_no or not issue_quantity:
                    st.error("❌ Please fill required fields")
                else:
                    try:
                        supabase_client.table('fuel_entries').insert({
                            'slip_no': slip_no,
                            'transaction_type': 'Issue',
                            'fuel_type': fuel_type,
                            'date': datetime.now().strftime("%Y-%m-%d"),
                            'vehicle_no': vehicle_no,
                            'vin_no': vin_no,
                            'registration_no': registration_no,
                            'model_no': model_no,
                            'allocated_to': allocated_to,
                            'opening_odometer': opening_odo,
                            'closing_odometer': closing_odo,
                            'km_run': km_run,
                            'issue_quantity': issue_quantity,
                            'mileage': mileage,
                            'recipient_name': recipient_name,
                            'receiver_name': receiver_name,
                            'approved_by': approved_by,
                            'remarks': remarks
                        }).execute()
                        st.success("✅ Issue saved successfully!")
                        st.balloons()
                        st.rerun()
                    except Exception as e:
                        if "duplicate key" in str(e).lower():
                            st.warning("⚠️ This Slip No already exists as an Issue. Try a different number.")
                        else:
                            st.error(f"Error: {str(e)}")
    
    # ============ REPORTS PAGE ============
    elif page == "📈 Reports":
        st.title("📈 Reports & Export")
        st.write("View, filter, and export transactions")
        st.divider()
        
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            fuel_filter = st.selectbox("Fuel Type", ["All", "Diesel", "Petrol"], key="ff_r")
        with col2:
            type_filter = st.selectbox("Transaction Type", ["All", "Purchase", "Issue"], key="tf_r")
        with col3:
            date_from = st.date_input("From Date", key="df_r")
        with col4:
            date_to = st.date_input("To Date", key="dt_r")
        with col5:
            if st.button("🔍 Filter", use_container_width=True, key="filter_r"):
                st.session_state.apply_filter = True
        
        if st.session_state.get('apply_filter', False) or st.button("📊 Load All Data"):
            try:
                entries = supabase_client.table('fuel_entries').select('*').execute()
                if entries.data:
                    df = pd.DataFrame(entries.data)
                    
                    if fuel_filter != "All":
                        df = df[df['fuel_type'] == fuel_filter]
                    if type_filter != "All":
                        df = df[df['transaction_type'] == type_filter]
                    
                    st.divider()
                    st.subheader(f"📋 Found {len(df)} records")
                    st.dataframe(df, use_container_width=True, hide_index=True)
                    
                    st.divider()
                    st.subheader("📥 Export Options")
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        excel_data = export_to_excel(entries.data)
                        st.download_button("📊 Excel", excel_data, "fuel_report.xlsx",
                                         "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                         use_container_width=True)
                    with col2:
                        pdf_data = export_to_pdf(entries.data)
                        st.download_button("📄 PDF", pdf_data, "fuel_report.pdf",
                                         "application/pdf", use_container_width=True)
                    with col3:
                        word_data = export_to_word(entries.data)
                        st.download_button("📝 Word", word_data, "fuel_report.docx",
                                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                         use_container_width=True)
                else:
                    st.info("📭 No entries found")
            except Exception as e:
                st.error(f"❌ Error: {e}")
    
    # ============ USERS PAGE ============
    elif page == "👥 Users":
        st.title("👥 Manage Users")
        st.divider()
        
        st.subheader("➕ Create New User")
        col1, col2, col3 = st.columns(3)
        with col1:
            new_username = st.text_input("Username", key="nu")
        with col2:
            new_password = st.text_input("Password", type="password", key="np")
        with col3:
            new_role = st.selectbox("Role", ["User", "Admin"], key="nr")
        
        if st.button("Create User", use_container_width=True, key="cu"):
            if not new_username or not new_password:
                st.error("❌ Please fill all fields")
            else:
                try:
                    supabase_client.table('users').insert({
                        'username': new_username,
                        'password': new_password,
                        'role': new_role
                    }).execute()
                    st.success("✅ User created!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Error: {e}")
        
        st.divider()
        st.subheader("👥 All Users")
        try:
            users = supabase_client.table('users').select('*').execute()
            if users.data:
                df = pd.DataFrame(users.data)
                st.dataframe(df[['username', 'role', 'created_at']], use_container_width=True, hide_index=True)
            else:
                st.info("📭 No users found")
        except Exception as e:
            st.error(f"❌ Error: {e}")
    
    # ============ PASSWORD PAGE ============
    elif page == "🔑 Password":
        st.title("🔑 Change Password")
        st.write("Update your login password")
        st.divider()
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            old_password = st.text_input("Old Password", type="password", key="op")
            new_password = st.text_input("New Password", type="password", key="np_c")
            confirm_password = st.text_input("Confirm Password", type="password", key="cp")
            
            if st.button("Update Password", use_container_width=True, key="up"):
                if new_password != confirm_password:
                    st.error("❌ Passwords do not match")
                else:
                    try:
                        response = supabase_client.table('users').select('*').eq('username', st.session_state.user['username']).execute()
                        if response.data and response.data[0]['password'] == old_password:
                            supabase_client.table('users').update({
                                'password': new_password
                            }).eq('id', st.session_state.user['id']).execute()
                            st.success("✅ Password changed!")
                        else:
                            st.error("❌ Old password incorrect")
                    except Exception as e:
                        st.error(f"❌ Error: {e}")